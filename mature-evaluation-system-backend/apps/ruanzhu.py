import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

# ====== 配置区 ======
# 项目代码根目录
PROJECT_DIR = r"D:\code\xiangmu_yang\mature-evaluation-system-frontend\src"

# 输出的 Word 文件
OUTPUT_DOCX = r"D:\code\xiangmu_yang\software_code2.docx"

# 要收集的代码文件后缀
INCLUDE_EXTS = {
    ".py", ".java", ".c", ".cpp", ".h", ".hpp",
    ".js", ".ts", ".vue", ".html", ".css",
    ".xml", ".json", ".yml", ".yaml", ".sh",
    ".php", ".go", ".rs", ".sql"
}

# 要忽略的目录
EXCLUDE_DIRS = {
    ".git", ".idea", ".vscode", "__pycache__",
    "node_modules", "dist", "build", "target",
    "out", "bin", "obj", ".venv", "venv"
}

# 要忽略的文件名
EXCLUDE_FILES = {
    "package-lock.json", "yarn.lock"
}

# 单个文件最大字符数，防止某些文件过大导致 Wordpip install python-docx 太卡
MAX_CHARS_PER_FILE = 500000
# ===================


def is_code_file(file_path: Path) -> bool:
    return file_path.suffix.lower() in INCLUDE_EXTS and file_path.name not in EXCLUDE_FILES


def should_skip_dir(dir_name: str) -> bool:
    return dir_name in EXCLUDE_DIRS


def collect_code_files(root_dir: Path):
    code_files = []
    for current_root, dirs, files in os.walk(root_dir):
        # 原地过滤目录
        dirs[:] = [d for d in dirs if not should_skip_dir(d)]

        for file_name in files:
            file_path = Path(current_root) / file_name
            if is_code_file(file_path):
                code_files.append(file_path)

    code_files.sort(key=lambda x: str(x).lower())
    return code_files


def set_document_style(doc: Document):
    # 正文样式
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    # 页面边距
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_code_to_doc(doc: Document, root_dir: Path, file_path: Path):
    relative_path = file_path.relative_to(root_dir)

    # 文件标题
    title = doc.add_paragraph()
    run = title.add_run(f"文件：{relative_path}")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)

    # 分隔说明
    info = doc.add_paragraph()
    info_run = info.add_run("代码内容如下：")
    info_run.font.name = "Times New Roman"
    info_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    info_run.font.size = Pt(10.5)

    # 读取代码
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = file_path.read_text(encoding="gbk")
        except UnicodeDecodeError:
            try:
                content = file_path.read_text(encoding="latin1")
            except Exception as e:
                content = f"[读取失败] {e}"

    # 限制长度
    if len(content) > MAX_CHARS_PER_FILE:
        content = content[:MAX_CHARS_PER_FILE] + "\n\n...[内容过长，已截断]"

    # 代码段
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(content)
    code_run.font.name = "Courier New"
    code_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_run.font.size = Pt(9)

    # 段落格式
    code_para.paragraph_format.space_before = Pt(0)
    code_para.paragraph_format.space_after = Pt(6)
    code_para.paragraph_format.left_indent = Pt(12)



def main():
    root_dir = Path(PROJECT_DIR)
    if not root_dir.exists():
        print(f"项目目录不存在：{PROJECT_DIR}")
        return

    files = collect_code_files(root_dir)
    if not files:
        print("没有找到符合条件的代码文件。")
        return

    doc = Document()
    set_document_style(doc)

    # 文档标题
    title = doc.add_paragraph()
    title_run = title.add_run("软件著作权申请 - 源代码文档")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_run.font.size = Pt(16)

    doc.add_paragraph(f"项目目录：{PROJECT_DIR}")
    doc.add_paragraph(f"共收集代码文件：{len(files)} 个")
    doc.add_page_break()

    for idx, file_path in enumerate(files, start=1):
        print(f"[{idx}/{len(files)}] 正在写入：{file_path}")
        add_code_to_doc(doc, root_dir, file_path)

    doc.save(OUTPUT_DOCX)
    print(f"\n完成，Word 已生成：{OUTPUT_DOCX}")


if __name__ == "__main__":
    main()