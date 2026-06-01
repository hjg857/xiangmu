"""
评估计分服务
实现五个维度的计分逻辑

计分公式说明（按老师要求）：
=================================
1. 观测点标准化: C_k = (C_i / S_i) × 5
   - C_i: 观测点实际得分
   - S_i: 观测点满分
   - C_k: 标准化后得分，满分5分

2. 二级指标得分: Σ(C_k × γ_k)
   - γ_k: 观测点权重（同一二级指标下的观测点权重和=1）
   - 满分5分

3. 一级指标得分: Σ(二级得分 × β_j)
   - β_j: 二级指标权重（同一一级指标下的二级权重和=1）
   - 满分5分

4. 总分: Σ(一级得分 × α_i)
   - α_i: 一级指标权重（五个维度权重和=1）
   - 满分5分
=================================
"""
from decimal import Decimal
from typing import Dict, List
import logging
from . import scoring_config as config

logger = logging.getLogger(__name__)


class ScoringService:
    """评估计分服务类"""
    
    def __init__(self, assessment):
        self.assessment = assessment
        self.school = assessment.school
        self.observation_scores = {}  # 观测点得分（5分制）
        self.secondary_scores = {}    # 二级指标得分（5分制）
        self.dimension_scores = {}    # 一级维度得分（5分制）

    def _count_selected_items(self, value) -> int:
        """
        统计多选题已选数量。
        前端一般传 list；如果旧数据传 dict 或字符串，也尽量兼容。
        """
        if value is None:
            return 0

        if isinstance(value, list):
            return len([item for item in value if item])

        if isinstance(value, dict):
            return len([k for k, v in value.items() if v])

        if isinstance(value, str):
            return 1 if value.strip() else 0

        return 0

    def calculate_all_scores(self) -> Dict[str, Decimal]:
        """
        计算所有维度得分
        返回: 各维度得分（5分制）和总分（5分制）
        """
        logger.info(f"开始计算评估 {self.assessment.id} 的得分")

        # 计算各维度得分（5分制）
        a_score = self.calculate_literacy_score()      # A 数据素养
        b_score = self.calculate_institution_score()   # B 数据制度
        c_score = self.calculate_behavior_score()      # C 数据行为
        d_score = self.calculate_asset_score()         # D 数据资产
        e_score = self.calculate_technology_score()    # E 数据技术

        # 存储一级维度得分
        self.dimension_scores = {
            'A': float(a_score),
            'B': float(b_score),
            'C': float(c_score),
            'D': float(d_score),
            'E': float(e_score),
        }

        # 计算总分（5分制）: Σ(一级得分 × α_i)
        total_score_5 = (
            float(a_score) * config.DIMENSION_WEIGHTS['A'] +
            float(b_score) * config.DIMENSION_WEIGHTS['B'] +
            float(c_score) * config.DIMENSION_WEIGHTS['C'] +
            float(d_score) * config.DIMENSION_WEIGHTS['D'] +
            float(e_score) * config.DIMENSION_WEIGHTS['E']
        )

        scores = {
            'literacy_score': Decimal(str(round(float(a_score), 4))),
            'institution_score': Decimal(str(round(float(b_score), 4))),
            'behavior_score': Decimal(str(round(float(c_score), 4))),
            'asset_score': Decimal(str(round(float(d_score), 4))),
            'technology_score': Decimal(str(round(float(e_score), 4))),
            'total_score': Decimal(str(round(total_score_5, 4))),
        }

        # 确定成熟度等级（基于5分制）
        scores['maturity_level'] = self.determine_maturity_level(total_score_5)

        logger.info(f"评估 {self.assessment.id} 计分完成:")
        logger.info(f"  A数据素养={a_score:.4f}, B数据制度={b_score:.4f}, C数据行为={c_score:.4f}")
        logger.info(f"  D数据资产={d_score:.4f}, E数据技术={e_score:.4f}")
        logger.info(f"  总分={total_score_5:.4f}（5分制）, 等级={scores['maturity_level']}")
        
        return scores

    def _normalize_score(self, raw_score: float, max_score: float) -> float:
        """
        观测点标准化: C_k = (C_i / S_i) × 5
        返回: 0-5分
        """
        if max_score <= 0:
            return 0.0
        normalized = (raw_score / max_score) * 5.0
        return min(normalized, 5.0)

    def _calculate_secondary_score(self, observation_scores: Dict[str, float], 
                                   observation_keys: List[str]) -> float:
        """
        计算二级指标得分: Σ(C_k × γ_k)
        observation_scores: {观测点编号: 5分制得分}
        observation_keys: 该二级指标包含的观测点编号列表
        返回: 0-5分
        """
        total = 0.0
        for key in observation_keys:
            score = observation_scores.get(key, 0.0)
            weight = config.OBSERVATION_WEIGHTS.get(key, 0.0)
            total += score * weight
        return total

    def _calculate_dimension_score(self, secondary_scores: Dict[str, float],
                                   secondary_keys: List[str]) -> float:
        """
        计算一级维度得分: Σ(二级得分 × β_j)
        secondary_scores: {二级指标编号: 5分制得分}
        secondary_keys: 该一级维度包含的二级指标编号列表
        返回: 0-5分
        """
        total = 0.0
        for key in secondary_keys:
            score = secondary_scores.get(key, 0.0)
            weight = config.SECONDARY_WEIGHTS.get(key, 0.0)
            total += score * weight
        return total

    def calculate_literacy_score(self) -> float:
        """
        计算A维度-数据素养得分（5分制）
        新版包含:
        A1 教师数据素养
        A2 学生数据素养
        """
        logger.info("计算数据素养得分")
        from apps.surveys.models import SurveyInstance

        teacher_instances = SurveyInstance.objects.filter(
            assessment=self.assessment,
            template__survey_type='teacher'
        )
        student_instances = SurveyInstance.objects.filter(
            assessment=self.assessment,
            template__survey_type='student'
        )

        # 计算教师、学生问卷观测点得分
        self._calc_literacy_observations('teacher', teacher_instances)
        self._calc_literacy_observations('student', student_instances)

        # A1 教师数据素养
        a1_score = self._calculate_secondary_score(
            self.observation_scores,
            ['A11', 'A12', 'A13', 'A14', 'A15']
        )

        # A2 学生数据素养
        # 注意：学生观测点编号使用正式权重表中的 A31-A35
        a2_score = self._calculate_secondary_score(
            self.observation_scores,
            ['A31', 'A32', 'A33', 'A34', 'A35']
        )

        self.secondary_scores['A1'] = a1_score
        self.secondary_scores['A2'] = a2_score

        dimension_score = self._calculate_dimension_score(
            self.secondary_scores,
            ['A1', 'A2']
        )

        logger.info(
            f"A维度: A1教师={a1_score:.4f}, A2学生={a2_score:.4f}, "
            f"维度得分={dimension_score:.4f}"
        )

        return dimension_score

    def _calc_literacy_observations(self, survey_type: str, instances) -> None:
        """计算数据素养问卷的观测点得分"""
        from apps.surveys.models import SurveyResponse
        
        survey_config = config.LITERACY_SURVEY_CONFIG.get(survey_type, {})
        
        responses = []
        for instance in instances:
            responses.extend(SurveyResponse.objects.filter(instance=instance))

        if not responses:
            logger.warning(f"没有{survey_type}问卷回答数据")
            for point_name in survey_config.keys():
                self.observation_scores[point_name] = 0.0
            return

        for point_name, point_config in survey_config.items():
            start, end = point_config['range']
            questions = list(range(start, end + 1))
            max_score = point_config['max_score']

            raw_score = self._calc_survey_raw_score(responses, questions)
            normalized = self._normalize_score(raw_score, max_score)
            self.observation_scores[point_name] = normalized
            
            logger.debug(f"{point_name}: 原始={raw_score:.2f}/{max_score}, 标准化={normalized:.4f}")

    def _calc_survey_raw_score(self, responses: List, question_nums: List[int]) -> float:
        """计算问卷题目的原始平均得分"""
        if not responses:
            return 0.0

        total_score = 0
        response_count = 0

        for response in responses:
            answers = response.answers or {}
            for q_num in question_nums:
                q_key = f'q{q_num}'
                if q_key in answers:
                    answer_value = answers[q_key]
                    if isinstance(answer_value, str):
                        total_score += config.SCALE_SCORE_MAPPING.get(answer_value, 0)
                    else:
                        total_score += int(answer_value)
            response_count += 1

        return total_score / response_count if response_count > 0 else 0.0

    def calculate_institution_score(self) -> float:
        """
        计算B维度-数据制度得分（5分制）
        包含: B1数据组织机构、B2数据人员配备、B3数据管理文件
        """
        logger.info("计算数据制度得分")

        try:
            inst = self.assessment.institution
        except:
            logger.warning("没有数据制度评估数据")
            return 0.0

        rules = config.INSTITUTION_SCORING_RULES

        # B11: 数据领导/工作小组
        b11_raw = rules['B11']['rules'].get(inst.leadership_group_type, 0)

        # 兼容旧数据：如果新字段为空，则使用旧布尔字段兜底
        if not inst.leadership_group_type:
            b11_raw = 10 if inst.has_leadership_group else 0

        self.observation_scores['B11'] = self._normalize_score(
            b11_raw,
            rules['B11']['max_score']
        )

        # B12: 数据组织运行情况
        b12_raw = self._apply_range_rules(inst.meeting_activity_count or 0, rules['B12']['rules'])
        self.observation_scores['B12'] = self._normalize_score(b12_raw, rules['B12']['max_score'])

        # B21: 数据专职/兼职管理人员
        if not inst.has_data_staff:
            b21_raw = 0
        else:
            staff_score = min((inst.fulltime_staff_count or 0) * 5 + (inst.parttime_staff_count or 0) * 3, 10)
            resp_score = 10 if inst.has_clear_responsibilities else 0
            b21_raw = staff_score + resp_score
        self.observation_scores['B21'] = self._normalize_score(b21_raw, rules['B21']['max_score'])

        # B22: 数据人员进修与培训
        if not inst.has_training:
            b22_raw = 0
        else:
            training_rules = [r for r in rules['B22']['sub_items'] if r.get('field') == 'training_count'][0]
            training_score = self._apply_range_rules(inst.training_count or 0, training_rules['rules'])
            cert_score = min((inst.national_cert_count or 0) * 5 + (inst.provincial_cert_count or 0) * 3 + (inst.city_cert_count or 0), 20)
            b22_raw = training_score + cert_score
        self.observation_scores['B22'] = self._normalize_score(b22_raw, rules['B22']['max_score'])

        # B31: 数据管理制度类文件
        # B31: 数据管理制度类文件
        management_status = inst.management_doc_status

        if management_status == 'clear_required':
            doc_count = inst.management_doc_count or 0
            doc_score = min(doc_count * 5, 20)
            quality_score = self._score_documents_with_llm(
                inst.management_doc_files,
                'management',
                20,
                inst
            )
            b31_raw = doc_score + quality_score

        elif management_status == 'follow_policy':
            b31_raw = 20

        elif management_status == 'self_awareness':
            b31_raw = 10

        else:
            # 兼容旧数据：如果新字段为空，则使用旧布尔字段兜底
            if not inst.has_management_doc:
                b31_raw = 0
            else:
                doc_count = inst.management_doc_count or 0
                doc_score = min(doc_count * 5, 20)
                quality_score = self._score_documents_with_llm(
                    inst.management_doc_files,
                    'management',
                    20,
                    inst
                )
                b31_raw = doc_score + quality_score

        self.observation_scores['B31'] = self._normalize_score(
            b31_raw,
            rules['B31']['max_score']
        )

        # B32: 数据实践指导类文件
        practice_status = inst.practice_doc_status

        if practice_status == 'published':
            doc_count = inst.practice_doc_count or 0
            doc_score = min(doc_count * 5, 20)
            quality_score = self._score_documents_with_llm(
                inst.practice_doc_files,
                'practice',
                20,
                inst
            )
            b32_raw = doc_score + quality_score

        elif practice_status == 'internal_training':
            b32_raw = 20

        elif practice_status == 'self_practice':
            b32_raw = 10

        else:
            # 兼容旧数据：如果新字段为空，则使用旧布尔字段兜底
            if not inst.has_practice_doc:
                b32_raw = 0
            else:
                doc_count = inst.practice_doc_count or 0
                doc_score = min(doc_count * 5, 20)
                quality_score = self._score_documents_with_llm(
                    inst.practice_doc_files,
                    'practice',
                    20,
                    inst
                )
                b32_raw = doc_score + quality_score

        self.observation_scores['B32'] = self._normalize_score(
            b32_raw,
            rules['B32']['max_score']
        )

        # 计算二级指标得分
        b1_score = self._calculate_secondary_score(self.observation_scores, ['B11', 'B12'])
        b2_score = self._calculate_secondary_score(self.observation_scores, ['B21', 'B22'])
        b3_score = self._calculate_secondary_score(self.observation_scores, ['B31', 'B32'])

        self.secondary_scores['B1'] = b1_score
        self.secondary_scores['B2'] = b2_score
        self.secondary_scores['B3'] = b3_score

        # 计算一级维度得分
        dimension_score = self._calculate_dimension_score(self.secondary_scores, ['B1', 'B2', 'B3'])

        logger.info(f"B维度: B1={b1_score:.4f}, B2={b2_score:.4f}, B3={b3_score:.4f}, 维度得分={dimension_score:.4f}")
        return dimension_score

    def calculate_behavior_score(self) -> float:
        """
        计算C维度-数据行为得分（5分制）
        包含: C1数据行为监测、C2数据应用成效
        """
        logger.info("计算数据行为得分")

        try:
            behavior = self.assessment.behavior
        except:
            logger.warning("没有数据行为评估数据")
            return 0.0

        rules = config.BEHAVIOR_SCORING_RULES

        # C11: 教师数据行为
        # C11: 教师数据行为
        c11_raw = 0

        # 1. 教师每周使用数字化设备开展教学的人均频次
        c11_raw += self._apply_range_rules(
            behavior.teacher_device_use_freq or 0,
            rules['C11']['sub_items'][0]['rules']
        )

        # 2. 教师每周使用数据相关平台的人均频次
        c11_raw += self._apply_range_rules(
            behavior.teacher_platform_use_freq or 0,
            rules['C11']['sub_items'][1]['rules']
        )

        # 3. 教师常态化开展的数据行为数量
        teacher_behavior_count = self._count_selected_items(
            behavior.teacher_data_behavior_items
        )
        c11_raw += self._apply_range_rules(
            teacher_behavior_count,
            rules['C11']['sub_items'][2]['rules']
        )

        self.observation_scores['C11'] = self._normalize_score(
            c11_raw,
            rules['C11']['max_score']
        )

        # C12: 学生数据行为
        c12_raw = 0

        # 1. 学生数字化学习设备配备情况
        student_device_rules = rules['C12']['sub_items'][0]['rules']
        c12_raw += student_device_rules.get(
            behavior.student_device_provision,
            0
        )

        # 2. 学生平台账号开通情况
        student_account_rules = rules['C12']['sub_items'][1]['rules']
        c12_raw += student_account_rules.get(
            behavior.student_account_status,
            0
        )

        # 3. 学生常态化实现的数据行为数量
        student_behavior_count = self._count_selected_items(
            behavior.student_data_behavior_items
        )
        c12_raw += self._apply_range_rules(
            student_behavior_count,
            rules['C12']['sub_items'][2]['rules']
        )

        self.observation_scores['C12'] = self._normalize_score(
            c12_raw,
            rules['C12']['max_score']
        )

        # C21: 数据应用特色成果
        c21_raw = 0
        for sub_item in rules['C21']['sub_items']:
            sub_score = sum((getattr(behavior, f, None) or 0) * s for f, s in sub_item['fields'].items())
            c21_raw += min(sub_score, sub_item['max_score'])
        self.observation_scores['C21'] = self._normalize_score(c21_raw, rules['C21']['max_score'])

        # C22: 数据应用社会影响
        # C22: 数据应用社会影响
        c22_raw = 0

        for sub_item in rules['C22']['sub_items']:
            if 'fields' in sub_item:
                sub_score = sum(
                    (getattr(behavior, field_name, None) or 0) * score
                    for field_name, score in sub_item['fields'].items()
                )
                c22_raw += min(sub_score, sub_item['max_score'])

            elif 'field' in sub_item:
                field_val = getattr(behavior, sub_item['field'], None) or 0

                if 'score_per_post' in sub_item:
                    c22_raw += min(
                        field_val * sub_item['score_per_post'],
                        sub_item['max_score']
                    )

                elif 'score_per_visit' in sub_item:
                    c22_raw += min(
                        field_val * sub_item['score_per_visit'],
                        sub_item['max_score']
                    )

        self.observation_scores['C22'] = self._normalize_score(
            c22_raw,
            rules['C22']['max_score']
        )

        # C23: 应用效果主观评价（来自问卷）
        c23_raw = self._calc_application_effect_score()
        self.observation_scores['C23'] = self._normalize_score(c23_raw, rules['C23']['max_score'])

        # 计算二级指标得分
        c1_score = self._calculate_secondary_score(self.observation_scores, ['C11', 'C12'])
        c2_score = self._calculate_secondary_score(self.observation_scores, ['C21', 'C22', 'C23'])

        self.secondary_scores['C1'] = c1_score
        self.secondary_scores['C2'] = c2_score

        # 计算一级维度得分
        dimension_score = self._calculate_dimension_score(self.secondary_scores, ['C1', 'C2'])

        logger.info(f"C维度: C1={c1_score:.4f}, C2={c2_score:.4f}, 维度得分={dimension_score:.4f}")
        return dimension_score

    def _calc_application_effect_score(self) -> float:
        """
        计算C23 教师对数据应用效果的主观评价得分。
        新版仅来自教师问卷 q49-q54，共6题，满分30分。
        """
        from apps.surveys.models import SurveyInstance, SurveyResponse

        c23_config = config.BEHAVIOR_SCORING_RULES['C23']
        start, end = c23_config['survey_ranges']['teacher']

        instances = SurveyInstance.objects.filter(
            assessment=self.assessment,
            template__survey_type='teacher'
        )

        responses = []
        for inst in instances:
            responses.extend(SurveyResponse.objects.filter(instance=inst))

        raw_score = self._calc_survey_raw_score(
            responses,
            list(range(start, end + 1))
        )

        # 记录教师C23的5分制得分，供前端图表或报告使用
        self.observation_scores['C23_teacher'] = self._normalize_score(
            raw_score,
            c23_config['max_score']
        )

        return raw_score

    def _get_data_volume_value(self, stat_method, volume) -> float:
        """
        根据统计方式返回有效数据量。
        unable：无法统计，计为0
        estimated/system_query：使用填写的数据量
        """
        if stat_method == 'unable':
            return 0.0

        try:
            return float(volume or 0)
        except (TypeError, ValueError):
            return 0.0

    def _calculate_total_data_volume(self, asset) -> float:
        """
        计算新版数据资产总量。
        包含：
        1. 教育教学数据
        2. 师生管理数据
        3. 数字资源数据
        4. 校园管理与行政数据
        5. 其他类型数据
        """
        total = 0.0

        total += self._get_data_volume_value(
            asset.teaching_data_stat_method,
            asset.teaching_data_volume
        )
        total += self._get_data_volume_value(
            asset.teacher_student_data_stat_method,
            asset.teacher_student_data_volume
        )
        total += self._get_data_volume_value(
            asset.digital_resource_data_stat_method,
            asset.digital_resource_data_volume
        )
        total += self._get_data_volume_value(
            asset.campus_admin_data_stat_method,
            asset.campus_admin_data_volume
        )

        # 其他类型数据为选填项，有值就计入
        try:
            total += float(asset.other_type_data_volume or 0)
        except (TypeError, ValueError):
            pass

        return total

    def calculate_asset_score(self) -> float:
        """
        计算D维度-数据资产得分（5分制）
        包含: D1数据资产意识、D2数据资产积累
        """
        logger.info("计算数据资产得分")

        try:
            asset = self.assessment.asset
            school = self.school
        except:
            logger.warning("没有数据资产评估数据")
            return 0.0

        from apps.surveys.models import SurveyInstance, SurveyResponse
        rules = config.ASSET_SCORING_RULES

        # 获取教师问卷回答
        teacher_instances = SurveyInstance.objects.filter(
            assessment=self.assessment,
            template__survey_type='teacher'
        )

        teacher_responses = []
        for inst in teacher_instances:
            teacher_responses.extend(SurveyResponse.objects.filter(instance=inst))

        # D11: 教师数据资产价值意识
        start, end = rules['D11']['survey_range']
        d11_raw = self._calc_survey_raw_score(
            teacher_responses,
            list(range(start, end + 1))
        )
        self.observation_scores['D11'] = self._normalize_score(
            d11_raw,
            rules['D11']['max_score']
        )

        # D12: 教师数据资产应用意识
        start, end = rules['D12']['survey_range']
        d12_raw = self._calc_survey_raw_score(
            teacher_responses,
            list(range(start, end + 1))
        )
        self.observation_scores['D12'] = self._normalize_score(
            d12_raw,
            rules['D12']['max_score']
        )

        # D13: 教师数据资产治理意识
        start, end = rules['D13']['survey_range']
        d13_raw = self._calc_survey_raw_score(
            teacher_responses,
            list(range(start, end + 1))
        )
        self.observation_scores['D13'] = self._normalize_score(
            d13_raw,
            rules['D13']['max_score']
        )

        # D21: 数据资产总量
        # 新规则：
        # 1. 未统一管理，D21 原始分直接给 fallback_raw_score
        # 2. 已统一管理但不能统计查询，D21 原始分直接给 fallback_raw_score
        # 3. 已统一管理且能够统计查询，则按数据总量区间计分
        can_calculate_asset_volume = (
                asset.has_unified_data_management is True
                and asset.can_query_data_assets is True
        )

        if not can_calculate_asset_volume:
            d21_raw = rules['D21'].get('fallback_raw_score', 4)
            total_volume = 0.0
        else:
            total_volume = self._calculate_total_data_volume(asset)
            d21_raw = self._apply_range_rules(
                total_volume,
                rules['D21']['rules']
            )

        self.observation_scores['D21'] = self._normalize_score(
            d21_raw,
            rules['D21']['max_score']
        )

        # D22: 人均数据资产量
        # 只有学校能够统计查询数据资产总量时，才计算人均数据资产量；
        # 前置条件不满足时，D22 直接计 0 分，避免 0GB 被区间规则误判为低档得分。
        if not can_calculate_asset_volume:
            d22_raw = 0
        else:
            total_people = (school.student_count or 0) + (school.teacher_count or 0)
            per_capita = total_volume / total_people if total_people > 0 else 0

            d22_raw = self._apply_range_rules(
                per_capita,
                rules['D22']['rules']
            )

        self.observation_scores['D22'] = self._normalize_score(
            d22_raw,
            rules['D22']['max_score']
        )

        # 计算二级指标得分
        d1_score = self._calculate_secondary_score(self.observation_scores, ['D11', 'D12', 'D13'])
        d2_score = self._calculate_secondary_score(self.observation_scores, ['D21', 'D22'])

        self.secondary_scores['D1'] = d1_score
        self.secondary_scores['D2'] = d2_score

        # 计算一级维度得分
        dimension_score = self._calculate_dimension_score(self.secondary_scores, ['D1', 'D2'])

        logger.info(f"D维度: D1={d1_score:.4f}, D2={d2_score:.4f}, 维度得分={dimension_score:.4f}")
        return dimension_score

    def calculate_technology_score(self) -> float:
        """
        计算E维度-数据技术得分（5分制）
        包含:
        E1 数据基础设施
        E2 数据安保水平
        """
        logger.info("计算数据技术得分")

        try:
            tech = self.assessment.technology
        except:
            logger.warning("没有数据技术评估数据")
            return 0.0

        rules = config.TECHNOLOGY_SCORING_RULES

        # E11: 数据硬件设施
        # 新规则：
        # 1. 未设立独立数据中心：数据中心标准项计0分
        # 2. 已设立独立数据中心：按完全达到/部分达到/未达到计分
        # 3. 生机比、师机比分别计分
        e11_raw = 0

        if tech.has_independent_data_center is True:
            data_center_value = tech.data_center_standard
            if data_center_value in rules['E11']['sub_items'][0]['rules']:
                e11_raw += rules['E11']['sub_items'][0]['rules'][data_center_value]
        else:
            # 未设立独立数据中心，该子项0分
            e11_raw += 0

        # 生机比
        student_ratio_rules = rules['E11']['sub_items'][1]['rules']
        if tech.student_device_ratio in student_ratio_rules:
            e11_raw += student_ratio_rules[tech.student_device_ratio]

        # 师机比
        teacher_ratio_rules = rules['E11']['sub_items'][2]['rules']
        if tech.teacher_device_ratio in teacher_ratio_rules:
            e11_raw += teacher_ratio_rules[tech.teacher_device_ratio]

        self.observation_scores['E11'] = self._normalize_score(
            e11_raw,
            rules['E11']['max_score']
        )

        # E12: 数据系统平台
        e12_raw = rules['E12']['rules'].get(tech.has_data_platform, 0)
        self.observation_scores['E12'] = self._normalize_score(
            e12_raw,
            rules['E12']['max_score']
        )

        # E21: 数据安全合规与认证
        # 平台建设管理模式：
        # self_built / mixed：继续按认证数量和认证比例计分
        # external：直接给5分制2.5分，也就是原始分10/20
        if tech.platform_build_mode == 'external':
            e21_raw = rules['E21'].get('external_platform_raw_score', 12)
            ratio_key = 'external'
            count_score = 0
            ratio_score = 0
        else:
            e21_raw = 0

            # 1. 数据安全认证数量
            count_rules = rules['E21']['sub_items'][0]['rules']
            count_score = self._apply_range_rules(
                tech.security_certified_count or 0,
                count_rules
            )
            e21_raw += count_score

            # 2. 数据安全认证比例
            ratio_rules = rules['E21']['sub_items'][1]['rules']
            ratio_value = tech.security_certified_ratio

            ratio_alias_map = {
                'zero': 'zero',
                '0': 'zero',
                'none': 'zero',

                'low': 'low',
                'below_40': 'low',
                'lt_40': 'low',
                '0_40': 'low',
                '0-40': 'low',
                '0<认定比例≤40%': 'low',
                '认定比例≤40%': 'low',

                'medium': 'medium',
                '40_80': 'medium',
                '40-80': 'medium',
                '40%<认定比例≤80%': 'medium',

                'high': 'high',
                'above_80': 'high',
                'gt_80': 'high',
                '认定比例>80%': 'high',
            }

            ratio_key = ratio_alias_map.get(ratio_value, ratio_value)
            ratio_score = ratio_rules.get(ratio_key, 0)
            e21_raw += ratio_score

        logger.info(
            f"E21计分: platform_build_mode={tech.platform_build_mode}, "
            f"security_certified_count={tech.security_certified_count}, "
            f"security_certified_ratio={tech.security_certified_ratio}, "
            f"ratio_key={ratio_key}, "
            f"count_score={count_score}, "
            f"ratio_score={ratio_score}, "
            f"e21_raw={e21_raw}"
        )

        self.observation_scores['E21'] = self._normalize_score(
            e21_raw,
            rules['E21']['max_score']
        )

        # E22: 数据风险事件记录
        # 新规则：发生过=0分，未发生=10分
        e22_raw = rules['E22']['rules'].get(tech.has_security_incident, 0)
        self.observation_scores['E22'] = self._normalize_score(
            e22_raw,
            rules['E22']['max_score']
        )

        # 计算二级指标得分
        e1_score = self._calculate_secondary_score(
            self.observation_scores,
            ['E11', 'E12']
        )
        e2_score = self._calculate_secondary_score(
            self.observation_scores,
            ['E21', 'E22']
        )

        self.secondary_scores['E1'] = e1_score
        self.secondary_scores['E2'] = e2_score

        # 计算一级维度得分
        dimension_score = self._calculate_dimension_score(
            self.secondary_scores,
            ['E1', 'E2']
        )

        logger.info(
            f"E维度: E1={e1_score:.4f}, E2={e2_score:.4f}, "
            f"维度得分={dimension_score:.4f}"
        )

        return dimension_score

    def _apply_range_rules(self, value: float, rules: List[Dict]) -> float:
        """应用范围规则计算得分"""
        for rule in rules:
            cond = rule['condition']
            if cond == 'lt' and value < rule['value']:
                return rule['score']
            elif cond == 'lte' and value <= rule['value']:
                return rule['score']
            elif cond == 'gt' and value > rule['value']:
                return rule['score']
            elif cond == 'gte' and value >= rule['value']:
                return rule['score']
            elif cond == 'between':
                min_val = rule.get('min', float('-inf'))
                max_val = rule.get('max', float('inf'))
                if min_val < value <= max_val:
                    return rule['score']
        return 0.0

    def determine_maturity_level(self, total_score: float) -> str:
        """
        根据最新规则确定等级：
        等级一：初始级 [0.0, 1.5]
        等级二：成长级 (1.5, 3.0]
        等级三：成熟级 (3.0, 4.0]
        等级四：创新级 (4.0, 5.0]
        """
        if total_score > 4.0:
            return 'leading'  # 对应“创新级/引领级”
        elif total_score > 3.0:
            return 'mature'  # 对应“成熟级”
        elif total_score > 1.5:
            return 'growing'  # 对应“成长级”
        else:
            return 'initial'  # 对应“初始级”

    def _score_documents_with_llm(self, doc_files: list, doc_type: str, 
                                   max_score: float = 20.0, institution=None) -> float:
        """使用大模型评分文档质量"""
        from apps.admin_panel.models import SystemConfig
        from .llm_service import LLMService

        # 检查缓存
        if institution:
            cached = getattr(institution, f'{doc_type}_doc_analysis', '') or ''
            if cached.strip():
                logger.info(f"使用缓存的{doc_type}文件分析结果")
                return max_score * 0.75 if doc_files else 0.0

        # 检查是否启用大模型
        llm_enabled = SystemConfig.get_config('llm_enabled', None)
        if llm_enabled is None:
            from django.conf import settings
            llm_enabled = bool(getattr(settings, 'DEEPSEEK_API_KEY', ''))

        if not llm_enabled or not doc_files:
            return max_score / 2.0 if doc_files else 0.0

        try:
            llm_service = LLMService()
            total_score = 0
            scored_count = 0
            all_analysis = []

            for doc_info in doc_files:
                try:
                    file_path = doc_info.get('path') or doc_info.get('url')
                    if not file_path:
                        continue

                    content = self._read_file_content(file_path)
                    if not content:
                        continue

                    if doc_type == 'management':
                        result = llm_service.score_management_document(content, max_score=max_score)
                    else:
                        result = llm_service.score_practice_document(content, max_score=max_score)

                    if result['success']:
                        total_score += result['score']
                        scored_count += 1
                        all_analysis.append(f"【{doc_info.get('name', '未知')}】\n{result.get('analysis', '')}")
                    else:
                        total_score += max_score / 2.0
                        scored_count += 1
                except Exception as e:
                    logger.error(f"评分文件出错: {e}")
                    total_score += max_score / 2.0
                    scored_count += 1

            # 保存分析结果
            if institution and all_analysis:
                try:
                    setattr(institution, f'{doc_type}_doc_analysis', '\n\n'.join(all_analysis))
                    institution.save(update_fields=[f'{doc_type}_doc_analysis'])
                except Exception as e:
                    logger.error(f"保存分析结果失败: {e}")

            return min(total_score / scored_count, max_score) if scored_count > 0 else max_score / 2.0

        except Exception as e:
            logger.error(f"大模型评分出错: {e}")
            return max_score / 2.0

    def _read_file_content(self, file_path: str) -> str:
        """读取文件内容"""
        from django.core.files.storage import default_storage
        import os

        try:
            if file_path.startswith('/media/'):
                file_path = file_path.replace('/media/', '')
            elif file_path.startswith('http'):
                from urllib.parse import urlparse
                file_path = urlparse(file_path).path.replace('/media/', '')

            if not default_storage.exists(file_path):
                return ""

            file_ext = os.path.splitext(file_path)[1].lower()

            if file_ext == '.pdf':
                return self._extract_pdf_content(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx_content(file_path)
            else:
                with default_storage.open(file_path, 'rb') as f:
                    content = self._decode_content(f.read(), file_path)

            return content[:8000] + "\n...(已截断)" if len(content) > 8000 else content

        except Exception as e:
            logger.error(f"读取文件失败: {e}")
            return ""

    def _decode_content(self, content_bytes: bytes, file_path: str) -> str:
        """尝试多种编码解码"""
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
            try:
                return content_bytes.decode(encoding)
            except:
                continue
        return ""

    def _extract_pdf_content(self, file_path: str) -> str:
        """提取PDF内容"""
        try:
            import PyPDF2
            from django.core.files.storage import default_storage
            with default_storage.open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return '\n'.join(page.extract_text() for page in reader.pages)
        except ImportError:
            return "PDF文件（需安装PyPDF2）"
        except Exception as e:
            return f"PDF提取失败: {e}"

    def _extract_docx_content(self, file_path: str) -> str:
        """提取DOCX内容"""
        try:
            import docx
            from django.core.files.storage import default_storage
            with default_storage.open(file_path, 'rb') as f:
                doc = docx.Document(f)
                text = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        text.append(' | '.join(c.text for c in row.cells))
                return '\n'.join(text)
        except ImportError:
            return "DOCX文件（需安装python-docx）"
        except Exception as e:
            return f"DOCX提取失败: {e}"
